from flask import Flask, render_template, request, jsonify, session, redirect, url_for, flash
import os
from werkzeug.utils import secure_filename
from OPENAI import call_gpt4o, call_gpt4o_test
import PyPDF2
import docx
import re
from OPENAI import get_risk_prompt_iteration_0, get_risk_prompt_iteration_1, format_output_with_highlights, analyze_risks_detailed,analyze_risks_initial
import json
from loguru import logger
import subprocess
import tempfile
import io
import sqlite3
from datetime import datetime
import pandas as pd
from datetime import timedelta
from urllib.parse import urlparse

from models import db, AnalysisRecord, User

app = Flask(__name__)

### For saving logic
# app.config['UPLOAD_FOLDER'] = 'uploads'
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///Analysis.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-me')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.permanent_session_lifetime = timedelta(days=7)


db.init_app(app)

with app.app_context():
    db.create_all()

@app.route("/save_analysis", methods=["POST"])
def save_analysis():
    data = request.get_json()

    if not data:
        return jsonify({"status": "error", "message": "No data received"}), 400

    try:
        record = AnalysisRecord(
            content_json=data["content_json"],
            analysis_type=data["analysis_type"],
            ownership=data["ownership"],
            risk_name=data["risk_name"],
            filename=data["filename"]
        )
        db.session.add(record)
        db.session.commit()
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_storage_obj):
    ext = file_storage_obj.filename.rsplit('.', 1)[1].lower()

    if ext == 'txt':
        return file_storage_obj.read().decode('utf-8')

    elif ext == 'pdf':
        reader = PyPDF2.PdfReader(file_storage_obj)
        return ' '.join(
            re.sub("\n", " ", page.extract_text()) 
            for page in reader.pages if page.extract_text()
        )

    elif ext == 'docx':
        # Reset stream pointer before reading
        file_storage_obj.seek(0)
        in_memory = io.BytesIO(file_storage_obj.read())
        doc = docx.Document(in_memory)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)

    return 'Unsupported file format.'

### This func for saving files.
# def extract_text(filepath):
#     ext = filepath.rsplit('.', 1)[1].lower()
#     if ext == 'txt':
#         with open(filepath, 'r', encoding='utf-8') as f:
#             return f.read()
#     elif ext == 'pdf':
#         with open(filepath, 'rb') as f:
#             reader = PyPDF2.PdfReader(f)
#             return ' '.join(re.sub("\n", " ", page.extract_text()) for page in reader.pages if page.extract_text())
#     elif ext == 'docx':
#         doc = docx.Document(filepath)
#         return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
#     return 'Unsupported file format.'

@app.route('/')
def index():
    return render_template("index.html")

@app.route('/risk_analysis', methods=['GET', 'POST'])
def upload_file():
    text = ''
    filename = ''
    analysis_result = []
    risk_lis = []
    risk_analysis_pairs = []
    selected_names = []
    analysis_dict = {}
    page = int(request.form.get("page", 0))

    try:
        analysis_dict = json.loads(request.form.get("analysis_dict", "{}"))
    except:
        analysis_dict = {}


    if request.method == 'POST':
        action = request.form.get('action')
        text = request.form.get('text', '')

        if action == 'upload':
            file = request.files.get('file')
            if file and allowed_file(file.filename):
                ### Uncover when saving uploaded file is needed
                # filename = secure_filename(file.filename)
                # filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                # file.save(filepath)
                # text = extract_text(filepath)
                filename = file.filename
                text = extract_text(file)

        elif action == 'analyze':
            if text:
                filename = request.form.get('filename', '')
                prompt = get_risk_prompt_iteration_0(text)
                analysis_result, risk_lis = analyze_risks_initial(call_gpt4o(prompt))
                # for i in range(len(analysis_result)):
                #     analysis_result[i] = format_output_with_highlights(analysis_result[i], "output5")
                risk_analysis_pairs = list(zip(risk_lis, analysis_result))

        elif action == 'reset':
            text = ''
            filename = ''
            analysis_result = []
            risk_lis = []
            risk_analysis_pairs = []
            selected_names = []
            analysis_dict = {}
            page = 0


    return render_template(
        'risk_analysis.html',
        text=text,
        # analysis_result=format_output_with_highlights(analysis_result, "output2"),
        risk_lis=risk_lis,
        selected_names=selected_names,
        risk_analysis_pairs=risk_analysis_pairs,
        analysis_dict=analysis_dict,
        current_page=page,
        filename=filename
    )

@app.route('/detailed_analysis', methods=['POST'])
def detailed_analysis():
    text = request.form.get('text', '')
    name = request.form.get('risk_name', '')

    if not name or not text:
        return "Invalid request", 400

    try:
        prompt = get_risk_prompt_iteration_1(
            text,
            level_0_risk=name.split("-")[0].strip(),
            level_1_risk=name.split("-")[1].strip()
        )
        result = call_gpt4o(prompt)
        detailed_result = analyze_risks_detailed(result)

    except Exception as e:
        logger.exception("Error during detailed analysis")
        return "Error during analysis", 500

    return render_template("detailed_analysis_partial.html", risk_name=name, detailed_result=detailed_result)

@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    data = request.json
    feedback = data.get('feedback')
    user_id = data.get('user_id')
    timestamp = datetime.utcnow().isoformat()

    conn = sqlite3.connect('./feedback.db')
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id TEXT,
            feedback TEXT,
            timestamp TEXT,
            project TEXT
        )
    ''')

    cursor.execute("""
        INSERT INTO feedback (user_id, feedback, timestamp, project)
        VALUES (?, ?, ?, ?)
    """, (user_id, feedback, timestamp, 'risk_web_system'))

    conn.commit()
    conn.close()

    return jsonify({'status': 'success'})

@app.route('/f')
def format():
    return render_template("format.html")
    # AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
    # return render_template("format.html", api_key=AZURE_OPENAI_API_KEY)

@app.route('/view_feedback')
def view_feedback():
    conn = sqlite3.connect('/home/feedback.db')
    cursor = conn.cursor()
    cursor.execute("SELECT user_id, feedback, timestamp, project FROM feedback ORDER BY timestamp DESC")
    feedback_data = cursor.fetchall()
    conn.close()
    return render_template('view_feedback.html', feedback_data=feedback_data)


DB_PATH = "Test.db" 

def parse_impact_level(impact_text):
    levels = re.findall(r'\b(High|Medium|Low)\b', impact_text)
    level_map = {'Low': 1, 'Medium': 2, 'High': 3}
    scores = [level_map[i] for i in levels if i in level_map]
    return sum(scores) / len(scores) if scores else None

def get_data(access_level=None):
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM Risks", conn)
    conn.close()

    # Full stats regardless of filter
    total_risks = df.copy()

    # Filtered data for selected level
    if access_level == "user":
        df = df[df["access_level"] == "user"]
    elif access_level == "group":
        df = df[df["access_level"] == "group"]

    # Risk type distribution
    risk_types = df["risk_name"].value_counts().reset_index()
    risk_types.columns = ["risk_name", "count"]

    # File-level average impact & likelihood
    df["impact_score"] = df["impact"].apply(parse_impact_level)
    likelihood_map = {'Low': 1, 'Medium': 2, 'High': 3}
    df["likelihood_score"] = df["likelihood"].map(likelihood_map)
    scatter_df = df.groupby("file_name").agg({
        "impact_score": "mean",
        "likelihood_score": "mean"
    }).dropna().reset_index()

    # Table data: filtered records
    table_data = df[["file_name", "risk_name", "likelihood", "impact", "access_level"]].to_dict(orient="records")

    # File count per risk type (in full dataset)
    file_count = total_risks.groupby("risk_name")["file_name"].nunique().reset_index()
    file_count.columns = ["risk_name", "file_count"]

    return {
        "risk_types": {
            "labels": risk_types["risk_name"].tolist(),
            "counts": risk_types["count"].tolist()
        },
        "scatter": {
            "file_names": scatter_df["file_name"].tolist(),
            "impact": scatter_df["impact_score"].tolist(),
            "likelihood": scatter_df["likelihood_score"].tolist()
        },
        "table": table_data,
        "file_risk_counts": {
            "labels": file_count["risk_name"].tolist(),
            "counts": file_count["file_count"].tolist()
        }
    }

@app.route("/risk_overview")
def risk_overview1():
    return render_template("risk_overview.html")

@app.route("/data/<level>")
def get_level_data(level):
    return jsonify(get_data(level))


# # --- Login ---
# @app.context_processor
# def inject_user():
#     """Make current user name available to all templates as `current_user_name`."""
#     return {
#         'current_user_name': session.get('user_name'),
#         'current_user_email': session.get('user_email')
#     }

# @app.route('/auth')
# def auth_page():
#     return render_template('auth.html')

# @app.post('/register')
# def register():
#     name = request.form.get('name', '').strip()
#     email = request.form.get('email', '').strip().lower()
#     password = request.form.get('password', '')

#     if not name or not email or not password:
#         flash('All fields are required.', 'danger')
#         return redirect(url_for('auth_page'))

#     if User.query.filter_by(email=email).first():
#         flash('Email is already registered.', 'warning')
#         return redirect(url_for('auth_page'))

#     u = User(name=name, email=email)
#     u.set_password(password)
#     db.session.add(u)
#     db.session.commit()

#     # Auto-login after registration
#     session.permanent = True
#     session['user_id'] = u.id
#     session['user_name'] = u.name
#     session['user_email'] = u.email

#     flash('Account created. Welcome!', 'success')
#     return redirect(url_for('index'))

# @app.post('/login')
# def login():
#     email = request.form.get('email', '').strip().lower()
#     password = request.form.get('password', '')

#     user = User.query.filter_by(email=email).first()
#     if not user or not user.check_password(password):
#         flash('Invalid email or password.', 'danger')
#         return redirect(url_for('auth_page'))

#     session.permanent = True
#     session['user_id'] = user.id
#     session['user_name'] = user.name
#     session['user_email'] = user.email

#     flash('Signed in successfully.', 'success')
#     return redirect(url_for('index'))

# @app.get('/logout')
# def logout():
#     session.clear()
#     flash('Signed out.', 'info')
#     return redirect(url_for('index'))

# @app.get('/switch_account')
# def switch_account():
#     session.clear()
#     return redirect(url_for('auth_page'))

ALLOWED_NEXT = {'/risk_analysis'}

@app.context_processor
def inject_user():
    """Make current user name available to all templates as `current_user_name`."""
    return {
        'current_user_name': session.get('user_name'),
        'current_user_email': session.get('user_email')
    }

def _get_safe_next():
    """从参数/表单里拿 next，并校验只允许回到白名单路径，避免开放重定向。"""
    next_url = request.args.get('next') or request.form.get('next') or ''
    if not next_url:
        return None
    try:
        path = urlparse(next_url).path
    except Exception:
        return None
    return next_url if path in ALLOWED_NEXT else None

def _redirect_after_auth(default_endpoint='index'):
    """登录/注册/登出后的统一跳转逻辑。"""
    nxt = _get_safe_next()
    if nxt:
        return redirect(nxt)
    return redirect(url_for(default_endpoint))

@app.route('/auth')
def auth_page():
    # 模板里会读取 request.args.get('next')，无需后端额外处理
    return render_template('auth.html')

@app.post('/register')
def register():
    name = request.form.get('name', '').strip()
    email = request.form.get('email', '').strip().lower()
    password = request.form.get('password', '')

    if not name or not email or not password:
        flash('All fields are required.', 'danger')
        return redirect(url_for('auth_page', next=request.form.get('next', '')))

    if User.query.filter_by(email=email).first():
        flash('Email is already registered.', 'warning')
        return redirect(url_for('auth_page', next=request.form.get('next', '')))

    u = User(name=name, email=email)
    u.set_password(password)
    db.session.add(u)
    db.session.commit()

    # Auto-login after registration
    session.permanent = True
    session['user_id'] = u.id
    session['user_name'] = u.name
    session['user_email'] = u.email

    flash('Account created. Welcome!', 'success')
    return _redirect_after_auth()

@app.post('/login')
def login():
    email = request.form.get('email', '').strip().lower()
    password = request.form.get('password', '')

    user = User.query.filter_by(email=email).first()
    if not user or not user.check_password(password):
        flash('Invalid email or password.', 'danger')
        return redirect(url_for('auth_page', next=request.form.get('next', '')))

    session.permanent = True
    session['user_id'] = user.id
    session['user_name'] = user.name
    session['user_email'] = user.email

    flash('Signed in successfully.', 'success')
    return _redirect_after_auth()

@app.get('/logout')
def logout():
    session.clear()
    flash('Signed out.', 'info')
    return _redirect_after_auth()

@app.get('/switch_account')
def switch_account():
    # 清会话后跳到 /auth；如果来源是白名单页，则把它作为 next 透传
    session.clear()
    ref = request.referrer or ''
    try:
        path = urlparse(ref).path
    except Exception:
        path = ''
    if path in ALLOWED_NEXT:
        return redirect(url_for('auth_page', next=ref))
    return redirect(url_for('auth_page'))

if __name__ == '__main__':
    host = os.environ.get('AZURE_WEBAPP_HOST', '0.0.0.0')
    port = int(os.environ.get('AZURE_WEBAPP_PORT', 8000))
    # Debug should be False in production (set via App Setting)
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    app.run(host=host, port=port, debug=debug)